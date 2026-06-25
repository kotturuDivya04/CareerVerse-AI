# =============================================================================
# modules/resume/reader.py  —  CareerVerse AI
# Resume-specific text extractor.
#
# Extracts plain text from PDF, DOCX, and TXT resume files and applies
# light post-processing suited for resume content:
#   - Strips lone page-number lines
#   - Normalises Unicode dashes, bullets and smart quotes to ASCII
#   - Collapses excessive blank lines
#   - Strips per-line leading/trailing whitespace
#
# Used by:
#   - modules/resume/resume_engine.py    (ATS scoring)
#   - modules/job_recommender/recommender_engine.py (skill extraction)
#   - modules/interview/question_generator.py      (project hint extraction)
#
# Intentionally self-contained — does NOT import from the plagiarism
# module so resume_engine can be tested independently of plagiarism deps.
# =============================================================================

from __future__ import annotations
import os
import re


# =============================================================================
# PUBLIC API
# =============================================================================

def extract_resume_text(filepath: str) -> str:
    """
    Extract and clean text from a resume file.

    Supported formats: .pdf  .docx  .txt
    Returns a plain-text string. Returns '' on any unrecoverable error
    so callers can check for an empty result and show a user-facing error.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        raw = _extract_pdf(filepath)
    elif ext == '.docx':
        raw = _extract_docx(filepath)
    elif ext == '.txt':
        raw = _extract_txt(filepath)
    else:
        return ''

    return _clean(raw)


# =============================================================================
# FORMAT EXTRACTORS
# =============================================================================

def _extract_pdf(filepath: str) -> str:
    """
    Extract text from all pages of a PDF using PyMuPDF (fitz).
    Pages joined with newline so section boundaries are preserved.
    """
    try:
        import fitz  # PyMuPDF — installed via requirements.txt
        doc   = fitz.open(filepath)
        pages = [page.get_text('text') for page in doc]
        doc.close()
        return '\n'.join(pages)
    except Exception as exc:
        print(f'[resume/reader] PDF extraction failed — {filepath}: {exc}')
        return ''


def _extract_docx(filepath: str) -> str:
    """
    Extract text from a DOCX file paragraph by paragraph.

    Table cells are also extracted (many resume templates use tables
    for their layout) to ensure skills and contact info are captured.
    """
    try:
        from docx import Document  # python-docx — installed via requirements.txt
        doc   = Document(filepath)
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Table cells (resume layouts often use tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in parts:
                        parts.append(text)

        return '\n'.join(parts)
    except Exception as exc:
        print(f'[resume/reader] DOCX extraction failed — {filepath}: {exc}')
        return ''


def _extract_txt(filepath: str) -> str:
    """
    Read a plain-text resume, trying UTF-8 then falling back to latin-1
    so accented characters (common in international student resumes)
    do not cause a crash.
    """
    for encoding in ('utf-8', 'latin-1', 'cp1252'):
        try:
            with open(filepath, 'r', encoding=encoding) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            print(f'[resume/reader] TXT extraction failed — {filepath}: {exc}')
            return ''
    return ''


# =============================================================================
# TEXT CLEANER
# =============================================================================

def _clean(text: str) -> str:
    """
    Post-process raw extracted text for resume analysis.

    Steps:
      1. Normalise Unicode punctuation to ASCII equivalents
      2. Strip leading/trailing whitespace per line
      3. Remove lines that are only digits (page numbers like "1", "2")
      4. Remove lines that are only special characters (e.g. "---", "•••")
      5. Collapse 3+ consecutive blank lines to a single blank line
    """
    if not text:
        return ''

    # ---- 1. Unicode normalisation ----
    replacements = {
        '\u2019': "'",  '\u2018': "'",    # smart single quotes
        '\u201c': '"',  '\u201d': '"',    # smart double quotes
        '\u2013': '-',  '\u2014': '-',    # en-dash, em-dash
        '\u2022': '-',  '\u00b7': '-',    # bullet, middle dot
        '\u2023': '-',  '\u25cf': '-',    # triangle bullet, black circle
        '\u00a0': ' ',                    # non-breaking space
        '\t':     ' ',                    # tabs to space
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # ---- 2 + 3 + 4. Per-line cleaning ----
    lines   = text.splitlines()
    cleaned = []
    for line in lines:
        line = line.strip()
        if not line:
            cleaned.append('')
            continue
        # Skip bare page numbers (1–3 digits alone on a line)
        if re.match(r'^\d{1,3}$', line):
            continue
        # Skip lines made entirely of punctuation/symbols (dividers)
        if re.match(r'^[\W_]+$', line):
            continue
        cleaned.append(line)

    text = '\n'.join(cleaned)

    # ---- 5. Collapse excessive blank lines ----
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()